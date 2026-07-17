"""Extração de texto de arquivos Word (.docx)."""


def extract_text_from_docx(caminho: str) -> str:
    from docx import Document

    doc = Document(caminho)
    partes = [p.text for p in doc.paragraphs if p.text.strip()]

    for tabela in doc.tables:
        for linha in tabela.rows:
            celulas = [c.text.strip() for c in linha.cells]
            if any(celulas):
                partes.append(" | ".join(celulas))

    return "\n".join(partes)
