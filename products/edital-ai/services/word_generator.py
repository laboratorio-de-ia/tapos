from pathlib import Path

from docx import Document
from docx.shared import Pt

from models import Edital, Analise


def gerar_word(edital: Edital, analise: Analise, caminho_saida: str) -> str:
    doc = Document()

    doc.add_heading("Análise de Edital — TAPOS edital-ai", level=1)

    tabela = doc.add_table(rows=0, cols=2)
    tabela.style = "Light Grid Accent 1"
    for campo, valor in [
        ("Número", edital.numero or "-"),
        ("Órgão", edital.orgao or "-"),
        ("Modalidade", edital.modalidade or "-"),
        ("Score de conformidade", f"{analise.score_conformidade:.0f}%"),
        ("Análise gerada por IA", "Sim" if analise.ia_utilizada else "Não (fallback determinístico)"),
    ]:
        linha = tabela.add_row().cells
        linha[0].text = campo
        linha[1].text = str(valor)

    doc.add_heading("Objeto", level=2)
    doc.add_paragraph(edital.objeto or "Não identificado.")

    doc.add_heading("Resumo Executivo", level=2)
    doc.add_paragraph(analise.resumo_executivo or "-")

    if analise.objetos_identificados:
        doc.add_heading("Objetos e Itens", level=2)
        t = doc.add_table(rows=1, cols=4)
        t.style = "Light List Accent 1"
        hdr = t.rows[0].cells
        hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = "Nº", "Descrição", "Quantidade", "Valor estimado"
        for o in analise.objetos_identificados:
            linha = t.add_row().cells
            linha[0].text = str(o.numero)
            linha[1].text = o.descricao
            linha[2].text = f"{o.quantidade or '-'} {o.unidade or ''}".strip()
            linha[3].text = f"R$ {o.valor_estimado:.2f}" if o.valor_estimado else "-"

    if analise.prazos_identificados:
        doc.add_heading("Prazos", level=2)
        for p in analise.prazos_identificados:
            doc.add_paragraph(f"{p.descricao}: {p.data_texto}", style="List Bullet")

    if analise.requisitos_identificados:
        doc.add_heading("Checklist de Requisitos de Habilitação", level=2)
        doc.add_paragraph(
            "Espaço reservado para anotações e conferência manual da equipe de compliance.",
        ).italic = True
        for r in analise.requisitos_identificados:
            doc.add_paragraph(f"☐ {r.descricao}", style="List Bullet")

    for titulo, itens in [
        ("Riscos", analise.riscos),
        ("Oportunidades", analise.oportunidades),
        ("Recomendações", analise.recomendacoes),
    ]:
        if itens:
            doc.add_heading(titulo, level=2)
            for item in itens:
                doc.add_paragraph(item, style="List Bullet")

    doc.add_paragraph()
    assinatura = doc.add_paragraph("Analisado por: _______________________    Data: ___/___/______")
    assinatura.runs[0].font.size = Pt(10)

    Path(caminho_saida).parent.mkdir(parents=True, exist_ok=True)
    doc.save(caminho_saida)
    return str(caminho_saida)
